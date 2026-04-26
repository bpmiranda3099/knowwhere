from __future__ import annotations

import argparse
import re
from pathlib import Path


def _parse_markdown_table(md: str) -> list[list[str]]:
    lines = [ln.rstrip("\n") for ln in md.splitlines() if ln.strip()]
    table_lines = [ln for ln in lines if ln.lstrip().startswith("|")]
    if len(table_lines) < 2:
        raise ValueError("No markdown table found.")

    # Drop the separator line (|---|---|...|)
    header = table_lines[0]
    sep = table_lines[1]
    if re.fullmatch(r"\|\s*-+\s*(\|\s*-+\s*)+\|", sep.strip()) is None:
        raise ValueError("Second row does not look like a markdown table separator.")

    data_lines = [header] + table_lines[2:]

    def split_row(row: str) -> list[str]:
        # Trim leading/trailing '|' then split
        parts = [p.strip() for p in row.strip().strip("|").split("|")]
        return parts

    rows = [split_row(r) for r in data_lines]
    col_count = len(rows[0])
    for r in rows:
        if len(r) != col_count:
            raise ValueError(f"Inconsistent column count in row: {r}")
    return rows


def _md_inline_to_plain_text(s: str) -> str:
    # Basic cleanup so the Word table is readable.
    s = s.replace("\\n", "\n")
    s = re.sub(r"`([^`]+)`", r"\1", s)  # inline code
    s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)  # bold
    s = re.sub(r"\s+", " ", s).strip()
    return s


def main() -> None:
    root = Path(__file__).resolve().parents[1]

    parser = argparse.ArgumentParser(
        description="Export a markdown table of unit test cases to a DOCX file."
    )
    parser.add_argument(
        "--in",
        dest="input_path",
        default=str(root / "docs" / "unit-test-cases.md"),
        help="Path to input markdown file (default: docs/unit-test-cases.md)",
    )
    parser.add_argument(
        "--out",
        dest="output_path",
        default=str(root / "docs" / "unit-test-cases.docx"),
        help="Path to output DOCX file (default: docs/unit-test-cases.docx)",
    )
    parser.add_argument(
        "--title",
        default="Unit Test Cases",
        help="DOCX heading title (default: Unit Test Cases)",
    )
    args = parser.parse_args()

    src = Path(args.input_path)
    if not src.is_absolute():
        src = (root / src).resolve()
    out = Path(args.output_path)
    if not out.is_absolute():
        out = (root / out).resolve()

    md = src.read_text(encoding="utf-8")
    rows = _parse_markdown_table(md)

    try:
        from docx import Document  # type: ignore
    except Exception as e:  # pragma: no cover
        raise SystemExit(
            "Missing dependency: python-docx. Install with: python -m pip install python-docx"
        ) from e

    doc = Document()
    doc.add_heading(args.title, level=1)

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"

    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.cell(i, j).text = _md_inline_to_plain_text(cell)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()

