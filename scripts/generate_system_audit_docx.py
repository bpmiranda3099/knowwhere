from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path


@dataclass(frozen=True)
class AuditRow:
    module_title: str
    criterion: str


def _default_rows() -> list[AuditRow]:
    return [
        AuditRow(
            "API Runtime (Server Setup & Defaults)",
            "\n".join(
                [
                    "Runs the HTTP server and exposes all API endpoints",
                    "Adds basic security protections for web requests",
                    "Controls which websites can call the API from a browser (allowed origins)",
                    "Applies request-rate limits with an allowlist for trusted sources",
                    "Supports deployments behind tunnels/gateways when configured",
                ]
            ),
        ),
        AuditRow(
            "API Access Control (API Key Guard)",
            "\n".join(
                [
                    "Optionally requires an access key for protected endpoints",
                    "Always allows health checks and readiness checks without a key",
                    "Allows the contact/feedback endpoint without a key (so the public UI can submit messages)",
                ]
            ),
        ),
        AuditRow(
            "API: POST /search",
            "\n".join(
                [
                    "Accepts a search question plus optional settings (result count, search style, and filters)",
                    "Supports different search styles: keyword-focused, meaning-focused, or a mix of both",
                    "Can return results as whole papers or as smaller passages",
                    "Can optionally re-order results for better relevance",
                    "Includes extra protection against heavy use (higher per-route request limit, but still bounded inputs)",
                ]
            ),
        ),
        AuditRow(
            "API: POST /embed",
            "\n".join(
                [
                    "Accepts a list of texts and returns a list of numeric representations (one per text)",
                    "Rejects empty inputs and enforces a minimum structure",
                    "Used to power meaning-based search and similarity comparisons",
                ]
            ),
        ),
        AuditRow(
            "API: POST /rerank",
            "\n".join(
                [
                    "Accepts a query and a list of candidate texts, then scores each candidate for relevance",
                    "Returns a score list aligned with the input order",
                    "If the reranking service is unavailable, safely returns an empty score list",
                ]
            ),
        ),
        AuditRow(
            "API: POST /ingest",
            "\n".join(
                [
                    "Starts an import job from a chosen source (arXiv, Crossref, or OpenAlex)",
                    "Accepts a query plus a target count (with a hard cap to prevent runaway imports)",
                    "Returns immediately with an “accepted” response and includes a process id for tracking",
                    "Strictly limited to prevent abuse (very low per-route rate limit)",
                ]
            ),
        ),
        AuditRow(
            "API: GET /health",
            "\n".join(
                [
                    "Provides a quick “is it running?” status for the API and supporting services",
                    "Can check specific dependencies on request (e.g., database, embedding service, reranker)",
                    "Returns a simple per-service status summary",
                ]
            ),
        ),
        AuditRow(
            "API: GET /ready",
            "\n".join(
                [
                    "Simple readiness endpoint for load balancers and deployments",
                    "Returns a “ready” status without doing dependency checks",
                ]
            ),
        ),
        AuditRow(
            "API: GET /stats",
            "\n".join(
                [
                    "Returns high-level counts for stored content (papers, passages, subjects, sources)",
                    "Used for sanity checks and reporting dataset size",
                ]
            ),
        ),
        AuditRow(
            "API: GET /logs",
            "\n".join(
                [
                    "Returns recent search activity records for review and troubleshooting",
                    "Supports a ‘limit’ parameter with a safe maximum to prevent huge responses",
                    "Returns logs sorted from newest to oldest",
                ]
            ),
        ),
        AuditRow(
            "API: POST /contact",
            "\n".join(
                [
                    "Accepts a message (name, email, role, message) and sends it to the configured recipient",
                    "Sends a confirmation copy back to the sender",
                    "If email is not configured, clearly reports that the feature is unavailable",
                    "Has a rate limit to reduce spam and abuse",
                ]
            ),
        ),
        AuditRow(
            "Core Search Logic (Search Service)",
            "\n".join(
                [
                    "Combines keyword matching with meaning-based matching to find relevant items",
                    "Produces a ranked list with relevance scores and short previews",
                    "Supports paper-level and passage-level results and common filters",
                ]
            ),
        ),
        AuditRow(
            "Search Mode: Keyword-focused (Lexical)",
            "\n".join(
                [
                    "Finds results by matching the user’s words to stored text fields",
                    "Prioritizes exact or close term matches when the wording matters",
                    "Works even when meaning-based services are unavailable",
                ]
            ),
        ),
        AuditRow(
            "Search Mode: Meaning-focused (Semantic)",
            "\n".join(
                [
                    "Finds results based on similarity of meaning, even when wording differs",
                    "Requires text-to-number conversion for the user query",
                    "Depends on stored meaning-representations being present for items",
                ]
            ),
        ),
        AuditRow(
            "Search Mode: Combined (Hybrid)",
            "\n".join(
                [
                    "Combines keyword and meaning signals into one overall relevance score",
                    "Uses adjustable weighting between keyword and meaning signals",
                    "Limits candidate pools for efficiency, then fuses and re-sorts results",
                ]
            ),
        ),
        AuditRow(
            "Search: Result Quality Heuristics",
            "\n".join(
                [
                    "Downranks low-information or ‘stub’ records so complete papers appear higher",
                    "Gives a small boost when both keyword and meaning signals agree",
                    "Removes duplicates (by DOI when available, otherwise by normalized title) and keeps the best-scoring entry",
                ]
            ),
        ),
        AuditRow(
            "Search: Filters (Year / Venue / Subject / Source)",
            "\n".join(
                [
                    "Lets users narrow results to a time range (from/to year)",
                    "Supports filtering by venue name, subject/category, and data source",
                    "Applies filters consistently across paper-level and passage-level search",
                ]
            ),
        ),
        AuditRow(
            "Search: Snippets / Previews",
            "\n".join(
                [
                    "Returns short previews of text so users can judge relevance quickly",
                    "Uses a fixed preview length for consistent UI display",
                ]
            ),
        ),
        AuditRow(
            "Search: Result Re-ordering + Relevance Gate",
            "\n".join(
                [
                    "Re-orders the top results using a second relevance check",
                    "Can be configured to return fewer results when confidence is low (strict relevance mode)",
                    "Falls back safely when strict mode would otherwise return nothing",
                ]
            ),
        ),
        AuditRow(
            "Meaning-based Matching (Embedding Service Integration)",
            "\n".join(
                [
                    "Converts text into a numeric form used for meaning-based comparisons",
                    "Must stay consistent with how the database stores these representations",
                    "Supports health checking to confirm the service is reachable",
                ]
            ),
        ),
        AuditRow(
            "Embedding Client: Request/Response Handling",
            "\n".join(
                [
                    "Sends text to the embedding service with a configured model name",
                    "Uses a timeout so requests don’t hang indefinitely",
                    "Accepts common response formats and rejects unknown formats clearly",
                ]
            ),
        ),
        AuditRow(
            "Relevance Improvement (Reranking Service Integration)",
            "\n".join(
                [
                    "Optionally re-orders the top candidates to improve relevance quality",
                    "Can be configured to be stricter about relevance (avoid weak matches)",
                    "Designed to degrade gracefully if the reranking service is unavailable",
                ]
            ),
        ),
        AuditRow(
            "Rerank Client: Timeouts and Test Bypass",
            "\n".join(
                [
                    "Uses a timeout so the system can recover from slow or stuck rerank requests",
                    "Allows reranking to be skipped during automated tests when the service isn’t running",
                    "Rejects missing configuration in normal operation to avoid silent failures",
                ]
            ),
        ),
        AuditRow(
            "Database (Stored Papers, Passages, and Search Data)",
            "\n".join(
                [
                    "Stores papers, optional passages, and the meaning-based representations used for search",
                    "Supports fast lookup for both keyword search and meaning-based search",
                    "Uses safe database access patterns to reduce risk of data errors or misuse",
                ]
            ),
        ),
        AuditRow(
            "Database Access Layer (Connection Pool)",
            "\n".join(
                [
                    "Maintains a reusable pool of database connections for performance",
                    "Applies connection timeouts and idle timeouts to avoid resource leaks",
                    "Provides a shared query function used across the system",
                ]
            ),
        ),
        AuditRow(
            "Database Setup (Schema / Migrations)",
            "\n".join(
                [
                    "Defines the database tables and relationships needed by the system",
                    "Sets up the database so new installations match the expected structure",
                    "Can be applied repeatedly without breaking an existing setup (safe re-run behavior)",
                ]
            ),
        ),
        AuditRow(
            "Configuration / Environment",
            "\n".join(
                [
                    "Central place to set how the system runs (addresses, ports, allowed websites, deployment mode)",
                    "Lets operators point the system to supporting services and choose which models to use",
                    "Enables operational controls like request limits and stricter relevance settings",
                ]
            ),
        ),
        AuditRow(
            "Configuration Validation (Startup Checks)",
            "\n".join(
                [
                    "Validates required settings at startup and stops early if something is missing",
                    "Ensures URLs and numbers are in the correct format",
                    "Enforces stricter requirements in non-test environments (e.g., reranking configured in production/dev)",
                ]
            ),
        ),
        AuditRow(
            "Data Ingestion (Pipeline Controller)",
            "\n".join(
                [
                    "Coordinates importing new research data into the system",
                    "Uses safe pacing and automatic retries to avoid overloading external services",
                    "Avoids duplicates and updates existing records when the same paper appears again",
                    "Supports long-running imports that can pause and continue later, with progress tracking",
                ]
            ),
        ),
        AuditRow(
            "Data Ingestion Sources (arXiv / Crossref / OpenAlex)",
            "\n".join(
                [
                    "Imports paper details from multiple external sources",
                    "Collects and standardizes key details (title, authors, year, identifiers, links) before saving",
                    "When available, brings in document text for better searching and previews",
                ]
            ),
        ),
        AuditRow(
            "Ingestion Utility: Retry + Backoff",
            "\n".join(
                [
                    "Retries temporary network failures automatically",
                    "Waits longer between retries to reduce repeated failures",
                    "Handles “too many requests” responses by waiting before retrying",
                ]
            ),
        ),
        AuditRow(
            "Ingestion Utility: Request Pacing (Rate Limit Helper)",
            "\n".join(
                [
                    "Adds intentional delays between requests to respect upstream limits",
                    "Prevents the system from overwhelming external data providers",
                ]
            ),
        ),
        AuditRow(
            "Ingestion Utility: Document Chunking",
            "\n".join(
                [
                    "Splits long text into overlapping passages for better search coverage",
                    "Controls maximum passage size and overlap to balance quality vs storage",
                ]
            ),
        ),
        AuditRow(
            "Ingestion Utility: PDF Text Extraction",
            "\n".join(
                [
                    "Downloads PDFs when available and extracts readable text",
                    "Fails safely (returns no text) when PDFs can’t be fetched or parsed",
                ]
            ),
        ),
        AuditRow(
            "Topic-based Import (Batch / Resume)",
            "\n".join(
                [
                    "Imports by topic lists to build a balanced dataset",
                    "Can distribute work across topics and sources to reach a target dataset size",
                    "Supports resume/checkpoints so long runs can be stopped and continued later",
                ]
            ),
        ),
        AuditRow(
            "Interactive Ingest Tool (CLI)",
            "\n".join(
                [
                    "Guides a user through starting an import (choose source, query, quantity)",
                    "Provides a topic-run mode with a saved progress file for resuming later",
                    "Prints progress updates (“heartbeats”) during long-running imports",
                ]
            ),
        ),
        AuditRow(
            "Topics Runner (Non-interactive CLI)",
            "\n".join(
                [
                    "Runs large topic imports from command-line flags (target, per-topic, sources, resume, dry-run)",
                    "Supports a dry-run mode that previews remaining work without importing",
                    "Tracks failures per topic and retries with increasing wait times",
                ]
            ),
        ),
        AuditRow(
            "Backfill Jobs",
            "\n".join(
                [
                    "Improves existing records after initial import (adds missing search-ready fields)",
                    "Fills in missing meaning-based representations for older items",
                    "Designed to be safe to run multiple times as data evolves",
                ]
            ),
        ),
        AuditRow(
            "Backfill: Search-ready Text Fields (TSV)",
            "\n".join(
                [
                    "Rebuilds the database’s search-ready text fields for papers and passages",
                    "Ensures keyword searching stays accurate after data changes",
                ]
            ),
        ),
        AuditRow(
            "Backfill: Missing Meaning Representations (Embeddings)",
            "\n".join(
                [
                    "Finds records missing meaning-representations and fills them in batches",
                    "Updates papers and passages separately",
                    "Uses a controlled batch size to avoid overloading the system",
                ]
            ),
        ),
        AuditRow(
            "Document Processing (PDF + Chunking Utilities)",
            "\n".join(
                [
                    "Extracts readable text from PDFs when available",
                    "Splits long documents into smaller passages for better searching and previews",
                    "Provides shared reliability tools used during import (retries and pacing)",
                ]
            ),
        ),
        AuditRow(
            "Contact / Email Delivery",
            "\n".join(
                [
                    "Allows users to send messages/feedback to the project team",
                    "Uses configurable email settings so it can work in different deployments",
                ]
            ),
        ),
        AuditRow(
            "Web UI (Static Playground + Docs)",
            "\n".join(
                [
                    "Provides a simple website for documentation and trying the search features",
                    "Includes a playground/test pages for sending example searches",
                    "Can be pointed at a different server address when needed (e.g., demo/public link)",
                    "Serves the website pages from standard paths (home and /web)",
                ]
            ),
        ),
        AuditRow(
            "Web Page: Landing Page (Home)",
            "\n".join(
                [
                    "Introduces the product and key features to new users",
                    "Provides navigation to the API docs and the search playground",
                    "Includes a contact form that sends messages to the backend contact endpoint",
                ]
            ),
        ),
        AuditRow(
            "Web Page: Search Playground",
            "\n".join(
                [
                    "Lets users run searches and compare result columns visually",
                    "Uses the proxied API path by default (so one public link works for both web and API)",
                    "Supports providing an API key in the UI when access control is enabled",
                ]
            ),
        ),
        AuditRow(
            "Web Page: API Documentation",
            "\n".join(
                [
                    "Explains request/response formats for each endpoint with examples",
                    "Documents how authentication and rate limits work",
                    "Serves as a static reference page (no login required)",
                ]
            ),
        ),
        AuditRow(
            "Reverse Proxy (Single Public Entry Point)",
            "\n".join(
                [
                    "Routes website ‘API’ requests to the backend automatically so users only need one public link",
                    "Keeps the website and API under one address for simpler demos and deployments",
                ]
            ),
        ),
        AuditRow(
            "Containerization / Local Deployment",
            "\n".join(
                [
                    "Runs the whole system locally in a consistent, repeatable way",
                    "Starts parts in the right order and checks readiness before use",
                    "Supports using pre-downloaded models to reduce setup time and avoid repeated downloads",
                ]
            ),
        ),
        AuditRow(
            "Docker Compose: Database Service",
            "\n".join(
                [
                    "Provides a local database for storing papers and search data",
                    "Keeps data across restarts using a persistent volume",
                    "Includes a health check so other services can wait until it’s ready",
                ]
            ),
        ),
        AuditRow(
            "Docker Compose: Embedding Service",
            "\n".join(
                [
                    "Runs the text-to-number service used for meaning-based search",
                    "Exposes a health endpoint to confirm readiness",
                    "Allows switching the model through configuration",
                ]
            ),
        ),
        AuditRow(
            "Docker Compose: Reranker Service",
            "\n".join(
                [
                    "Runs the result re-ordering service used to improve relevance quality",
                    "Exposes a health endpoint to confirm readiness",
                    "Allows switching the model through configuration",
                ]
            ),
        ),
        AuditRow(
            "Docker Compose: API Service",
            "\n".join(
                [
                    "Runs the backend that connects the web UI, database, and model services",
                    "Configured to use internal service addresses when running in Docker",
                    "Includes a health check for automated startup and monitoring",
                ]
            ),
        ),
        AuditRow(
            "Docker Compose: Schema Setup Job",
            "\n".join(
                [
                    "Applies the database setup script automatically",
                    "Waits for the database health check before running",
                    "Designed to be safe even if run multiple times",
                ]
            ),
        ),
        AuditRow(
            "Docker Compose: Web Service",
            "\n".join(
                [
                    "Serves the static website and proxies API requests under the same address",
                    "Supports live updates to static files via a mounted folder during development",
                ]
            ),
        ),
        AuditRow(
            "Testing & CI Signals",
            "\n".join(
                [
                    "Small tests to validate key functions and rules",
                    "System-level tests to confirm components work together as expected",
                    "Browser-based checks to ensure the website and playground behave correctly",
                    "Quality checks for supporting ML services, including coverage targets",
                ]
            ),
        ),
        AuditRow(
            "Observability / Operations",
            "\n".join(
                [
                    "Provides simple “is it running?” checks for each service",
                    "Exposes operational information useful for monitoring and troubleshooting (where enabled)",
                    "Includes safe default settings that reduce risk during public demos or heavy use",
                ]
            ),
        ),
    ]


def generate_docx(out_path: Path, title: str, rows: list[AuditRow]) -> None:
    try:
        from docx import Document  # type: ignore
    except Exception as e:  # pragma: no cover
        raise RuntimeError(
            "Missing dependency 'python-docx'. Install with: python3 -m pip install python-docx"
        ) from e

    document = Document()
    document.add_heading(title, level=1)

    meta = document.add_paragraph()
    meta.add_run("Generated: ").bold = True
    meta.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))

    document.add_paragraph(
        "Table fields capture modules and their primary software criteria "
        "(Characteristic / functionality / Specific Feature)."
    )

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Module Title"
    hdr_cells[1].text = "Software Criterion(Characteristic / functionality / Specific Feature )"

    for r in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = r.module_title
        row_cells[1].text = r.criterion

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    out_path = repo_root / "docs" / "system-audit.docx"
    generate_docx(
        out_path=out_path,
        title="KnowWhere System Audit",
        rows=_default_rows(),
    )
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()

